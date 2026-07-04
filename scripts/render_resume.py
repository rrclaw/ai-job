#!/usr/bin/env python3
"""Render a markdown resume to PDF (headless Chrome) and/or docx (python-docx).

Usage:
    python3 render_resume.py resume.md --pdf --docx [--css ../templates/resume.css] [--out-dir .]

Markdown subset supported: h1-h3, bullet lists, bold, links, plain paragraphs,
horizontal rules. Keep resumes inside this subset — fancier markdown means the
docx branch silently degrades.
"""
import argparse
import html
import pathlib
import re
import subprocess
import sys

CHROME = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"


def md_to_html_body(md):
    out, in_ul = [], False

    def close_ul():
        nonlocal in_ul
        if in_ul:
            out.append("</ul>")
            in_ul = False

    def inline(s):
        s = html.escape(s, quote=False)
        # re-allow the one raw tag resumes use: <span class="date">...</span>
        s = re.sub(r'&lt;span class="date"&gt;(.*?)&lt;/span&gt;',
                   r'<span class="date">\1</span>', s)
        s = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", s)
        s = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', s)
        return s

    for line in md.splitlines():
        if re.match(r"\s*$", line):
            close_ul()
            continue
        m = re.match(r"(#{1,3})\s+(.*)", line)
        if m:
            close_ul()
            n = len(m.group(1))
            out.append(f"<h{n}>{inline(m.group(2))}</h{n}>")
            continue
        if re.match(r"\s*(---|\*\*\*)\s*$", line):
            close_ul()
            out.append("<hr>")
            continue
        m = re.match(r"\s*[-*•]\s+(.*)", line)
        if m:
            if not in_ul:
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{inline(m.group(1))}</li>")
            continue
        close_ul()
        out.append(f"<p>{inline(line.strip())}</p>")
    close_ul()
    return "\n".join(out)


def to_pdf(md_path, css_path, out_dir):
    md = md_path.read_text(encoding="utf-8")
    css = css_path.read_text(encoding="utf-8") if css_path.exists() else ""
    html_doc = (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<style>{css}</style></head><body>{md_to_html_body(md)}</body></html>"
    )
    html_file = out_dir / (md_path.stem + ".render.html")
    pdf_file = out_dir / (md_path.stem + ".pdf")
    html_file.write_text(html_doc, encoding="utf-8")
    subprocess.run(
        [CHROME, "--headless", "--disable-gpu", "--no-pdf-header-footer",
         f"--print-to-pdf={pdf_file}", html_file.as_uri()],
        check=True, capture_output=True, timeout=120,
    )
    print(f"pdf: {pdf_file}")
    return pdf_file


def to_docx(md_path, out_dir):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")

    def add_runs(par, text):
        pos = 0
        for m in re.finditer(r"\*\*([^*]+)\*\*|\[([^\]]+)\]\(([^)]+)\)", text):
            if m.start() > pos:
                par.add_run(text[pos:m.start()])
            if m.group(1):
                par.add_run(m.group(1)).bold = True
            else:
                par.add_run(f"{m.group(2)} ({m.group(3)})")
            pos = m.end()
        if pos < len(text):
            par.add_run(text[pos:])

    for line in md_path.read_text(encoding="utf-8").splitlines():
        line = re.sub(r'\s*<span class="date">(.*?)</span>', r"（\1）", line)
        if re.match(r"\s*$", line) or re.match(r"\s*(---|\*\*\*)\s*$", line):
            continue
        m = re.match(r"(#{1,3})\s+(.*)", line)
        if m:
            h = doc.add_heading(level=len(m.group(1)))
            add_runs(h, m.group(2))
            if len(m.group(1)) == 1:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        m = re.match(r"\s*[-*•]\s+(.*)", line)
        if m:
            add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
            continue
        add_runs(doc.add_paragraph(), line.strip())

    out = out_dir / (md_path.stem + ".docx")
    doc.save(out)
    print(f"docx: {out}")
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("resume")
    ap.add_argument("--pdf", action="store_true")
    ap.add_argument("--docx", action="store_true")
    ap.add_argument("--css", default=str(pathlib.Path(__file__).parent.parent / "templates" / "resume.css"))
    ap.add_argument("--out-dir", default=None)
    args = ap.parse_args()

    md_path = pathlib.Path(args.resume).resolve()
    out_dir = pathlib.Path(args.out_dir).resolve() if args.out_dir else md_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    if not (args.pdf or args.docx):
        args.pdf = args.docx = True
    if args.pdf:
        to_pdf(md_path, pathlib.Path(args.css), out_dir)
    if args.docx:
        to_docx(md_path, out_dir)


if __name__ == "__main__":
    main()
